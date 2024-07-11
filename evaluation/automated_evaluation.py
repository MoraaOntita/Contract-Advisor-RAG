import os
import json
from docx import Document
from evaluation_tools import EvaluationTools
from rag_pipeline import RAGSystem

# Paths to documents
contract_path = "/home/moraa/Documents/10_academy/week-11/data/evaluate/Robinson Advisory.docx"
qa_path = "/home/moraa/Documents/10_academy/week-11/data/evaluate/Robinson Q&A.docx"

def load_document(file_path):
    doc = Document(file_path)
    text = [p.text for p in doc.paragraphs if p.text.strip() != ""]
    return text

def load_qa(file_path):
    doc = Document(file_path)
    qa_pairs = []
    for para in doc.paragraphs:
        if para.text.strip() != "":
            qa_pairs.append(para.text.split("? "))
    return qa_pairs

def main():
    # Load the documents
    contract_text = load_document(contract_path)
    qa_pairs = load_qa(qa_path)

    # Initialize RAG system
    rag_system = RAGSystem()

    # Initialize EvaluationTools
    evaluation_tools = EvaluationTools(rag_system)

    # Prepare test queries and expected responses
    test_queries = [qa[0] for qa in qa_pairs]
    expected_responses = [qa[1] for qa in qa_pairs]

    # Run evaluation
    metrics = evaluation_tools.evaluate(test_queries, expected_responses)

    # Save results
    evaluation_tools.save_results(metrics, output_file="evaluation/results.json")

    # Print results
    print("Evaluation results:")
    for metric, value in metrics.items():
        print(f"{metric}: {value}")

if __name__ == "__main__":
    main()
