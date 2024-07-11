import random
from docx import Document
from rag_pipeline import RAGSystem

# Paths to documents
contract_path = "/home/moraa/Documents/10_academy/week-11/data/evaluate/Robinson Advisory.docx"
qa_path = "/home/moraa/Documents/10_academy/week-11/data/evaluate/Robinson Q&A.docx"

def load_document(file_path):
    doc = Document(file_path)
    text = [p.text for p in doc.paragraphs if p.text.strip() != ""]
    return text

def load_qa(file_path):
    doc = Document(file_path)
    qa_pairs = []
    for para in doc.paragraphs:
        if para.text.strip() != "":
            qa_pairs.append(para.text.split("? "))
    return qa_pairs

def save_for_manual_review(responses, output_file):
    with open(output_file, "w") as f:
        for query, response in responses:
            f.write(f"Question: {query}\n")
            f.write(f"Generated Response: {response}\n")
            f.write("\n")

def main():
    # Load the documents
    contract_text = load_document(contract_path)
    qa_pairs = load_qa(qa_path)

    # Initialize RAG system
    rag_system = RAGSystem()

    # Prepare test queries
    test_queries = [qa[0] for qa in qa_pairs]

    # Select a subset for manual review
    sample_queries = random.sample(test_queries, min(5, len(test_queries)))

    # Generate responses
    responses = [(query, rag_system.query(query)) for query in sample_queries]

    # Save for manual review
    save_for_manual_review(responses, "evaluation/manual_review.txt")

    print("Sample responses saved for manual review in 'evaluation/manual_review.txt'.")

if __name__ == "__main__":
    main()
