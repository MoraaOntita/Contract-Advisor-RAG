import json
import sys
import os
import time
from sklearn.metrics import precision_score, recall_score, f1_score
from nltk.translate.bleu_score import sentence_bleu
from rouge_score import rouge_scorer
import matplotlib.pyplot as plt
from docx import Document
import config

# Add the src directory to the Python path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'src')))

from rag_pipeline import RAGSystem  # Ensure this matches your actual module structure

class EvaluationTools:
    def __init__(self, rag_system):
        self.rag_system = rag_system

    def evaluate(self, test_queries, expected_responses):
        start_time = time.time()
        generated_responses = [self.rag_system.query(query) for query in test_queries]
        response_time = time.time() - start_time

        # Accuracy and Relevance (simplified)
        accuracy = sum([gen == exp for gen, exp in zip(generated_responses, expected_responses)]) / len(test_queries)

        # Precision, Recall, F1 Score
        precision = precision_score(expected_responses, generated_responses, average='micro', zero_division=0)
        recall = recall_score(expected_responses, generated_responses, average='micro', zero_division=0)
        f1 = f1_score(expected_responses, generated_responses, average='micro', zero_division=0)

        # BLEU and ROUGE
        bleu_scores = [sentence_bleu([exp], gen) for gen, exp in zip(generated_responses, expected_responses)]
        bleu_score_avg = sum(bleu_scores) / len(bleu_scores)
        
        scorer = rouge_scorer.RougeScorer(['rouge1', 'rougeL'], use_stemmer=True)
        rouge_scores = [scorer.score(exp, gen) for gen, exp in zip(generated_responses, expected_responses)]
        rouge1_avg = sum([score['rouge1'].fmeasure for score in rouge_scores]) / len(rouge_scores)
        rougeL_avg = sum([score['rougeL'].fmeasure for score in rouge_scores]) / len(rouge_scores)

        return {
            "accuracy": accuracy,
            "response_time": response_time,
            "precision": precision,
            "recall": recall,
            "f1_score": f1,
            "bleu_score": bleu_score_avg,
            "rouge1_score": rouge1_avg,
            "rougeL_score": rougeL_avg
        }

    def save_results(self, metrics, output_file='evaluation/results.json'):
        with open(output_file, 'w') as f:
            json.dump(metrics, f)

    def plot_results(self, metrics):
        plt.figure(figsize=(10, 5))
        plt.bar(metrics.keys(), metrics.values())
        plt.title("RAG System Evaluation Metrics")
        plt.show()

def load_docx(file_path):
    doc = Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

def load_questions_answers(file_path):
    doc = Document(file_path)
    questions = []
    answers = []
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            questions.append(para.text)
        elif para.style.name == 'Normal':
            answers.append(para.text)
    return questions, answers

if __name__ == "__main__":
    # Load contract
    contract_text = load_docx(config.CONTRACT_PATH)
    print("Contract Loaded:\n", contract_text)

    # Load questions and answers
    questions, answers = load_questions_answers(config.QA_PATH)
    print("Questions and Answers Loaded:\n", questions, answers)

    # Initialize the RAG system
    rag_system = RAGSystem()

    # Initialize the evaluation tools
    evaluation_tools = EvaluationTools(rag_system)

    # Run evaluation
    metrics = evaluation_tools.evaluate(questions, answers)

    # Print results
    for metric, value in metrics.items():
        print(f"{metric}: {value}")

    # Save results to a JSON file
    evaluation_tools.save_results(metrics)

    # Plot results
    evaluation_tools.plot_results(metrics)
